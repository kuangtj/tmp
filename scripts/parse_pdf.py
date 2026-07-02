#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Parse paper PDF and supplementary files into raw_* tables.

Design goals for the new IPM pipeline:
- Keep every page image traceable for downstream targeted structure resolution.
- Insert semantic text / figure / table blocks from DotsOCR JSON.
- Register supplementary PDFs, tables, text, images, structure files and sequence files.
- Preserve rich provenance in metadata_json / table_json.

Outputs:
- raw_text_block
- raw_asset
- raw_figure
- raw_table
- data/work/{doc_id}/pages/page_XXX.png
- data/work/{doc_id}/figures/
- data/work/{doc_id}/tables/
- data/work/{doc_id}/supplementary/
"""

import argparse
import json
import re
import shutil
import subprocess
import sys
import time
import uuid
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import fitz  # PyMuPDF
import pandas as pd
from PIL import Image
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from ipm_eagle.db.sqlite import get_conn


# -----------------------------------------------------------------------------
# Basic helpers
# -----------------------------------------------------------------------------

def uid(prefix: str, *parts: Any) -> str:
    key = "|".join(map(str, parts))
    return f"{prefix}_{uuid.uuid5(uuid.NAMESPACE_URL, key).hex[:16]}"


def jdump(x: Any) -> str:
    return json.dumps(x if x is not None else {}, ensure_ascii=False)


def clean(x: Any) -> str:
    return re.sub(r"\s+", " ", str(x or "").strip())


def safe_name(x: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", str(x or "")).strip("_") or "item"


def read_text_safe(path: Path, max_chars: Optional[int] = None) -> str:
    path = Path(path)
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
    except Exception:
        text = ""
    return text[:max_chars] if max_chars else text


def load_json_safe(path: Path) -> Any:
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except UnicodeDecodeError:
        return json.loads(Path(path).read_text(encoding="utf-8-sig", errors="ignore"))
    except Exception:
        return {}


def json_safe_value(v: Any) -> Any:
    if pd.isna(v):
        return ""
    if hasattr(v, "isoformat"):
        try:
            return v.isoformat()
        except Exception:
            pass
    return v.item() if hasattr(v, "item") else v


def df_to_table_json(df: pd.DataFrame, source: str, meta: Dict[str, Any]) -> Dict[str, Any]:
    df2 = df.copy()
    df2.columns = [clean(c) or f"col_{i+1}" for i, c in enumerate(df2.columns)]
    rows = []
    for _, row in df2.iterrows():
        rows.append({str(k): json_safe_value(v) for k, v in row.to_dict().items()})
    return {
        "source": source,
        "columns": [str(c) for c in df2.columns],
        "rows": rows,
        "n_rows": int(df2.shape[0]),
        "n_cols": int(df2.shape[1]),
        "metadata": meta,
    }


# -----------------------------------------------------------------------------
# DB insert helpers
# -----------------------------------------------------------------------------

def insert_text(conn, block_id: str, doc_id: str, page_no: Optional[int], section: str, text: str, meta: Optional[Dict[str, Any]] = None) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO raw_text_block
        (block_id, doc_id, page_no, section, text, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (block_id, doc_id, page_no, section, text or "", jdump(meta or {})),
    )


def insert_asset(conn, asset_id: str, doc_id: str, asset_type: str, page_no: Optional[int], figure_ref: str, table_ref: str, file_path: Path, meta: Optional[Dict[str, Any]] = None) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO raw_asset
        (asset_id, doc_id, asset_type, page_no, figure_ref, table_ref, file_path, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (asset_id, doc_id, asset_type, page_no, figure_ref or "", table_ref or "", str(file_path), jdump(meta or {})),
    )


def insert_figure(conn, figure_id: str, doc_id: str, page_no: Optional[int], figure_ref: str, file_path: Path, caption: str, meta: Optional[Dict[str, Any]] = None) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO raw_figure
        (figure_id, doc_id, page_no, figure_ref, file_path, caption, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (figure_id, doc_id, page_no, figure_ref or "", str(file_path), caption or "", jdump(meta or {})),
    )


def insert_table(conn, table_id: str, doc_id: str, page_no: Optional[int], table_ref: str, file_path: Path, table_json: Optional[Dict[str, Any]] = None, meta: Optional[Dict[str, Any]] = None) -> None:
    conn.execute(
        """
        INSERT OR REPLACE INTO raw_table
        (table_id, doc_id, page_no, table_ref, file_path, table_json, metadata_json)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (table_id, doc_id, page_no, table_ref or "", str(file_path), jdump(table_json or {}), jdump(meta or {})),
    )


def update_document_status(conn, doc_id: str, status: str, meta_patch: Optional[Dict[str, Any]] = None) -> None:
    row = conn.execute("SELECT metadata_json FROM raw_document WHERE doc_id=?", (doc_id,)).fetchone()
    old_meta = {}
    if row and row["metadata_json"]:
        try:
            old_meta = json.loads(row["metadata_json"])
        except Exception:
            old_meta = {}
    if meta_patch:
        old_meta.update(meta_patch)
    conn.execute(
        "UPDATE raw_document SET status=?, metadata_json=?, updated_at=CURRENT_TIMESTAMP WHERE doc_id=?",
        (status, jdump(old_meta), doc_id),
    )


def clear_existing_raw_outputs(conn, doc_id: str) -> None:
    """Delete parser outputs for one document before re-parse."""
    for table in ["planned_tasks", "raw_text_block", "raw_asset", "raw_figure", "raw_table"]:
        conn.execute(f"DELETE FROM {table} WHERE doc_id=?", (doc_id,))


# -----------------------------------------------------------------------------
# DotsOCR parsing helpers
# -----------------------------------------------------------------------------

TEXT_CATEGORIES = {
    "title", "section-header", "text", "caption", "footnote", "formula", "list-item",
}
DROP_TEXT_CATEGORIES = {"page-header", "page-footer", "page-number", "page_number", "watermark"}
FIGURE_CATEGORIES = {"figure", "picture", "image"}
TABLE_CATEGORIES = {"table"}

TEXT_ASSET_SECTIONS = {
    "Title", "Section-header", "Text", "Caption",
    "supplementary_Title", "supplementary_Section-header", "supplementary_Text", "supplementary_Caption",
    "title", "section-header", "text", "caption",
    "supplementary_title", "supplementary_section-header", "supplementary_text", "supplementary_caption",
}


def norm_category(cat: Any) -> str:
    c = clean(cat).lower().replace(" ", "-").replace("_", "-")
    return c


def get_text(x: Any) -> str:
    if not isinstance(x, dict):
        return ""
    vals = []
    for k in ["text", "content", "markdown", "html", "latex"]:
        v = x.get(k)
        if isinstance(v, str) and v.strip():
            vals.append(v.strip())
    return "\n".join(vals)


def is_text_block_category(cat: Any) -> bool:
    c = norm_category(cat)
    return c not in DROP_TEXT_CATEGORIES and c in TEXT_CATEGORIES


def is_caption_block(cat: Any, text: str) -> bool:
    c = norm_category(cat)
    if c == "caption":
        return True
    return bool(re.match(r"^\s*(Fig\.|Figure|Scheme|Table)\s*\d+", clean(text), re.I))


def is_figure_block_category(cat: Any) -> bool:
    return norm_category(cat) in FIGURE_CATEGORIES


def is_table_block_category(cat: Any) -> bool:
    return norm_category(cat) in TABLE_CATEGORIES


def ref_from_caption(text: str) -> str:
    m = re.match(r"^\s*((?:Fig\.|Figure|Scheme|Table)\s*(?:S)?\d+[A-Za-z]?)", text or "", re.I)
    return m.group(1) if m else ""


def bbox_xyxy(bbox: Any, w: int, h: int) -> Optional[List[int]]:
    if bbox is None:
        return None
    vals = None
    if isinstance(bbox, dict):
        vals = [bbox.get("x0", bbox.get("left")), bbox.get("y0", bbox.get("top")), bbox.get("x1", bbox.get("right")), bbox.get("y1", bbox.get("bottom"))]
    elif isinstance(bbox, (list, tuple)):
        if len(bbox) == 4 and all(isinstance(x, (int, float)) for x in bbox):
            vals = list(bbox)
        elif bbox and isinstance(bbox[0], (list, tuple)):
            xs = [p[0] for p in bbox if len(p) >= 2]
            ys = [p[1] for p in bbox if len(p) >= 2]
            vals = [min(xs), min(ys), max(xs), max(ys)] if xs and ys else None
    if not vals or any(v is None for v in vals):
        return None
    x0, y0, x1, y1 = map(float, vals)
    if max(x0, y0, x1, y1) <= 1.5:
        x0, x1 = x0 * w, x1 * w
        y0, y1 = y0 * h, y1 * h
    x0 = max(0, min(w - 1, int(x0)))
    y0 = max(0, min(h - 1, int(y0)))
    x1 = max(0, min(w, int(x1)))
    y1 = max(0, min(h, int(y1)))
    if x1 <= x0 or y1 <= y0:
        return None
    return [x0, y0, x1, y1]


def bbox_area_ratio(xyxy: Optional[List[int]], width: int, height: int) -> float:
    if not xyxy or width <= 0 or height <= 0:
        return 0.0
    x0, y0, x1, y1 = xyxy
    return max(0, x1 - x0) * max(0, y1 - y0) / float(width * height)


def valid_crop_bbox(xyxy: Optional[List[int]], width: int, height: int, min_area_ratio: float = 0.0005, max_area_ratio: float = 0.90) -> bool:
    r = bbox_area_ratio(xyxy, width, height)
    return min_area_ratio <= r <= max_area_ratio


def crop(img_path: Path, bbox: Any, out_path: Path) -> Optional[List[int]]:
    im = Image.open(img_path).convert("RGB")
    xyxy = bbox_xyxy(bbox, im.width, im.height)
    if not xyxy:
        return None
    out_path.parent.mkdir(parents=True, exist_ok=True)
    im.crop(xyxy).save(out_path)
    return xyxy


def flatten_cells(obj: Any) -> List[Dict[str, Any]]:
    cells: List[Dict[str, Any]] = []

    def walk(x: Any) -> None:
        if isinstance(x, dict):
            cat = x.get("category") or x.get("type") or x.get("label") or x.get("class")
            bbox = x.get("bbox") or x.get("box") or x.get("position") or x.get("float_xyxy")
            text = get_text(x)
            if cat and (bbox is not None or text):
                cells.append({"category": str(cat), "bbox": bbox, "text": text, "raw": x})
            for v in x.values():
                if isinstance(v, (dict, list)):
                    walk(v)
        elif isinstance(x, list):
            for v in x:
                walk(v)

    walk(obj)
    return cells


def extract_page_no_from_json_path(path: Path, fallback: int) -> int:
    name = path.name
    for pat in [r"page[_-]?(\d+)", r"_(\d+)\.json$", r"(\d+)\.json$"]:
        m = re.search(pat, name, re.I)
        if m:
            return int(m.group(1))
    return fallback


def load_cells_from_json_files(json_paths: List[Path]) -> List[Dict[str, Any]]:
    cells: List[Dict[str, Any]] = []
    for j, p in enumerate(sorted(json_paths), start=0):
        obj = load_json_safe(p)
        page_no = extract_page_no_from_json_path(p, j)
        for c in flatten_cells(obj):
            c["source_json_path"] = str(p)
            c["source_json_index"] = j
            c["page_no"] = page_no
            cells.append(c)
    return cells


def merge_md_files(md_paths: List[Path]) -> str:
    parts = []
    for i, p in enumerate(sorted(md_paths), start=1):
        text = read_text_safe(p)
        if text.strip():
            parts.append(f"\n\n<!-- DOTSOCR_MD_SOURCE index={i} path={p} -->\n\n{text}")
    return "\n\n".join(parts).strip()


def run_dotsocr(input_path: Path, out_dir: Path, dots_root: str, dots_python: str, num_thread: int, port: int) -> Tuple[List[Path], List[Path], Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    dots_root_p = Path(dots_root).expanduser().resolve()
    input_path = Path(input_path).expanduser().resolve()
    before = {p.resolve() for p in out_dir.rglob("*") if p.suffix.lower() in [".json", ".md"]}
    t0 = time.time()
    cmd = [
        dots_python, 
        str(dots_root_p / "dots_ocr" / "parser.py"),
        str(input_path),
        "--num_thread", str(num_thread),
        "--port", str(port),
        "--model_name", "dotsocr",
        "--output", str(out_dir.resolve()),
    ]
    log_path = out_dir / f"{input_path.stem}.log"
    try:
        with log_path.open("w", encoding="utf-8") as f:
            subprocess.run(cmd, cwd=str(dots_root_p), stdout=f, stderr=subprocess.STDOUT, check=True)
    except subprocess.CalledProcessError as e:
        tail = read_text_safe(log_path)[-5000:]
        raise RuntimeError(
            f"DotsOCR failed for {input_path}\n"
            f"log_path={log_path}\n"
            f"command={' '.join(map(str, cmd))}\n"
            f"---- log tail ----\n{tail}"
        ) from e

    after = {p.resolve() for p in out_dir.rglob("*") if p.suffix.lower() in [".json", ".md"]}
    new_files = list(after - before)
    json_files = [p for p in new_files if p.suffix.lower() == ".json"]
    md_files = [p for p in new_files if p.suffix.lower() == ".md"]
    if not json_files:
        json_files = [p for p in out_dir.rglob("*.json") if p.stat().st_mtime >= t0 - 1]
    if not md_files:
        md_files = [p for p in out_dir.rglob("*.md") if p.stat().st_mtime >= t0 - 1]
    if not json_files:
        raise RuntimeError(f"DotsOCR json output not found for {input_path}; log_path={log_path}")
    return sorted(json_files), sorted(md_files), log_path


def render_pdf_pages(pdf_path: Path, out_dir: Path, doc_id: str, conn, prefix: str = "page", asset_type: str = "page_image", source_meta: Optional[Dict[str, Any]] = None, register_asset: bool = True, dpi: int = 200) -> List[Tuple[int, Path]]:
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(str(pdf_path))
    pages = []
    source_meta = source_meta or {}
    for i, page in enumerate(doc, start=0):
        img_path = out_dir / f"{prefix}_{i:03d}.png"
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        pix.save(str(img_path))
        pages.append((i, img_path))
        if register_asset:
            asset_id = uid("page", doc_id, asset_type, str(pdf_path), i)
            insert_asset(
                conn, asset_id, doc_id, asset_type, i, "", "", img_path,
                {
                    "source": "pymupdf_page_render",
                    "pdf": str(pdf_path),
                    "page_no": i,
                    "dpi": dpi,
                    "purpose": "targeted_resolution_and_planning",
                    **source_meta,
                },
            )
    return pages


def bind_captions(objects: List[Dict[str, Any]], captions: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    for cap in captions:
        ref = (cap.get("ref") or "").lower()
        if not ref:
            continue
        target_kind = "table" if ref.startswith("table") else "figure"
        cbox = cap.get("xyxy")
        if not cbox:
            continue
        cx = (cbox[0] + cbox[2]) / 2
        cy = (cbox[1] + cbox[3]) / 2
        best = None
        best_score = 1e18
        for obj in objects:
            if obj.get("kind") != target_kind or obj.get("page_no") != cap.get("page_no"):
                continue
            obox = obj.get("xyxy")
            if not obox:
                continue
            ox = (obox[0] + obox[2]) / 2
            oy = (obox[1] + obox[3]) / 2
            score = abs(cx - ox) * 0.2 + abs(cy - oy)
            if score < best_score:
                best = obj
                best_score = score
        if best:
            best["caption"] = cap.get("text", "")
            cap["bound_object_id"] = best.get("id")
            cap["bound_object_ref"] = best.get("ref")
            cap["bound_object_type"] = best.get("kind")
    return objects, captions


def write_page_text_assets(conn, doc_id: str, page_text_map: Dict[int, List[Dict[str, str]]], out_dir: Path, asset_type: str, source_pdf: Path, min_chars: int = 80) -> int:
    out_dir.mkdir(parents=True, exist_ok=True)
    n_assets = 0
    for page_no in sorted(page_text_map):
        parts = []
        block_ids = []
        sections = []
        for b in page_text_map[page_no]:
            text = clean(b.get("text"))
            section = clean(b.get("section"))
            block_id = clean(b.get("block_id"))
            if text:
                block_ids.append(block_id)
                sections.append(section)
                parts.append(f"[{section}]\n{text}")
        page_text = "\n\n".join(parts).strip()
        if len(page_text) < min_chars:
            continue
        file_path = out_dir / f"page_{int(page_no):03d}_text.md"
        file_path.write_text(page_text, encoding="utf-8")
        asset_id = uid("text_asset", doc_id, asset_type, source_pdf, page_no)
        insert_asset(
            conn, asset_id, doc_id, asset_type, page_no, "", "", file_path,
            {
                "source": "raw_text_block_page_aggregation",
                "source_pdf": str(source_pdf),
                "page_no": page_no,
                "block_ids": block_ids,
                "sections": sorted(set(sections)),
                "n_blocks": len(block_ids),
                "n_chars": len(page_text),
                "ingest_policy": "page_level_text_asset",
            },
        )
        n_assets += 1
    return n_assets


# -----------------------------------------------------------------------------
# PDF parsing
# -----------------------------------------------------------------------------

def parse_pdf_with_dotsocr_once(
    doc_id: str,
    pdf_path: Path,
    work_dir: Path,
    conn,
    dots_root: str,
    dots_python: str,
    num_thread: int,
    dots_port: int,
    kind: str = "main",
    register_page_images: bool = True,
) -> Dict[str, Any]:
    pdf_path = Path(pdf_path)
    safe = safe_name(pdf_path.stem)

    if kind == "main":
        page_dir = work_dir / "pages"
        dots_dir = work_dir / "dotsocr"
        fig_dir = work_dir / "figures"
        table_dir = work_dir / "tables"
        text_dir = work_dir / "text"
        page_prefix = "page"
        source_meta = {"document_part": "main"}
        page_asset_type = "page_image"
        fig_asset_type = "figure_image"
        table_asset_type = "table_image"
        text_asset_type = "text_page"
        text_prefix = "text"
        section_prefix = ""
        fig_id_prefix = "fig"
        table_id_prefix = "table"
        caption_id_prefix = "caption"
        ocr_input = pdf_path
    else:
        for_copy = work_dir
        for_copy.mkdir(parents=True, exist_ok=True)
        dst_pdf = for_copy / pdf_path.name
        if pdf_path.resolve() != dst_pdf.resolve():
            shutil.copy2(pdf_path, dst_pdf)
        pdf_path = dst_pdf
        pdf_asset_id = uid("supp_pdf", doc_id, pdf_path.name)
        insert_asset(
            conn, pdf_asset_id, doc_id, "supplementary_pdf", None, "", "", pdf_path,
            {
                "source_file": str(pdf_path),
                "parsed_by": "dotsocr_json_blocks",
                "document_part": "supplementary",
                "note": "original supplementary pdf; semantic blocks inserted separately",
            },
        )
        page_dir = work_dir / f"{safe}_pages"
        dots_dir = work_dir / f"{safe}_dotsocr"
        fig_dir = work_dir / f"{safe}_figures"
        table_dir = work_dir / f"{safe}_tables"
        text_dir = work_dir / f"{safe}_text"
        page_prefix = f"{safe}_page"
        source_meta = {"document_part": "supplementary", "supplementary_name": pdf_path.name}
        page_asset_type = "supplementary_page_image"
        fig_asset_type = "supplementary_figure_image"
        table_asset_type = "supplementary_table_image"
        text_asset_type = "supplementary_text_page"
        text_prefix = "supp_text"
        section_prefix = "supplementary_"
        fig_id_prefix = "supp_fig"
        table_id_prefix = "supp_table"
        caption_id_prefix = "supp_caption"
        ocr_input = pdf_path

    for d in [page_dir, dots_dir, fig_dir, table_dir, text_dir]:
        d.mkdir(parents=True, exist_ok=True)

    pages = render_pdf_pages(
        pdf_path=pdf_path,
        out_dir=page_dir,
        doc_id=doc_id,
        conn=conn,
        prefix=page_prefix,
        asset_type=page_asset_type,
        source_meta=source_meta,
        register_asset=register_page_images,
    )
    page_img_map = {int(page_no): img_path for page_no, img_path in pages}

    json_paths, md_paths, log_path = run_dotsocr(ocr_input, dots_dir / "pdf", dots_root, dots_python, num_thread, dots_port)
    full_md = merge_md_files(md_paths)
    full_md_path = text_dir / "full_text.md"
    full_md_path.write_text(full_md, encoding="utf-8")

    cells = load_cells_from_json_files(json_paths)
    (text_dir / "dotsocr_sources.json").write_text(
        json.dumps(
            {
                "num_md_files": len(md_paths),
                "num_json_files": len(json_paths),
                "md_paths": [str(x) for x in md_paths],
                "json_paths": [str(x) for x in json_paths],
                "merged_md_path": str(full_md_path),
                "num_cells": len(cells),
                "document_part": kind,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    captions: List[Dict[str, Any]] = []
    objects: List[Dict[str, Any]] = []
    page_text_map: Dict[int, List[Dict[str, str]]] = {}
    n_text = 0
    n_skipped = 0

    for i, cell in enumerate(cells):
        page_no = cell.get("page_no")
        if isinstance(page_no, str) and page_no.isdigit():
            page_no = int(page_no)

        cat = clean(cell.get("category"))
        text = clean(cell.get("text"))
        bbox = cell.get("bbox")
        norm_cat = norm_category(cat)
        page_img = page_img_map.get(int(page_no)) if page_no is not None else None

        xyxy = None
        im_w = 0
        im_h = 0
        if page_img and bbox is not None:
            try:
                im = Image.open(page_img)
                im_w, im_h = im.width, im.height
                xyxy = bbox_xyxy(bbox, im_w, im_h)
            except Exception:
                xyxy = None

        meta = {
            "source": "dotsocr_json_block",
            "pdf": str(ocr_input),
            "document_part": kind,
            "page_image": str(page_img) if page_img else "",
            "dotsocr_json": cell.get("source_json_path") or "",
            "dotsocr_json_index": cell.get("source_json_index"),
            "all_dotsocr_json": [str(x) for x in json_paths],
            "dotsocr_log": str(log_path),
            "category": cat,
            "norm_category": norm_cat,
            "bbox": bbox,
            "xyxy": xyxy,
            "page_no_from_dotsocr": page_no,
            "raw_block": cell.get("raw"),
        }

        # Text-like blocks.
        if text and is_text_block_category(cat):
            n_text += 1
            section_name = f"{section_prefix}{norm_cat}"
            block_id = uid(text_prefix, doc_id, pdf_path.name, page_no, i, norm_cat, text[:80])
            insert_text(conn, block_id, doc_id, page_no, section_name, text, meta)
            if page_no is not None:
                page_text_map.setdefault(int(page_no), []).append({"block_id": block_id, "section": section_name, "text": text})
            if is_caption_block(cat, text):
                captions.append({
                    "caption_id": uid(caption_id_prefix, doc_id, pdf_path.name, page_no, i),
                    "page_no": page_no,
                    "ref": ref_from_caption(text),
                    "text": text,
                    "bbox": bbox,
                    "xyxy": xyxy,
                })
            continue

        # Caption text whose category may not be exactly caption.
        if text and is_caption_block(cat, text):
            n_text += 1
            section_name = f"{section_prefix}caption"
            block_id = uid(text_prefix, doc_id, pdf_path.name, page_no, i, "caption", text[:80])
            insert_text(conn, block_id, doc_id, page_no, section_name, text, meta)
            if page_no is not None:
                page_text_map.setdefault(int(page_no), []).append({"block_id": block_id, "section": section_name, "text": text})
            captions.append({
                "caption_id": uid(caption_id_prefix, doc_id, pdf_path.name, page_no, i),
                "page_no": page_no,
                "ref": ref_from_caption(text),
                "text": text,
                "bbox": bbox,
                "xyxy": xyxy,
            })
            continue

        # Figure/table blocks need image bbox.
        if not page_img or bbox is None or not xyxy or not valid_crop_bbox(xyxy, im_w, im_h):
            n_skipped += 1
            continue

        if is_table_block_category(cat):
            count = len([x for x in objects if x["kind"] == "table" and x["page_no"] == page_no]) + 1
            table_ref = f"{safe}_page{int(page_no or 0):03d}_table{count:03d}" if kind != "main" else f"page{int(page_no or 0):03d}_table{count:03d}"
            img_out = table_dir / f"{table_ref}.png"
            xyxy2 = crop(page_img, bbox, img_out)
            if xyxy2:
                table_id = uid(table_id_prefix, doc_id, pdf_path.name, page_no, i)
                json_out = table_dir / f"{table_ref}.json"
                json_out.write_text(json.dumps(cell.get("raw"), ensure_ascii=False, indent=2), encoding="utf-8")
                objects.append({
                    "kind": "table", "id": table_id, "ref": table_ref, "page_no": page_no,
                    "file_path": img_out, "json_path": json_out, "text": text,
                    "bbox": bbox, "xyxy": xyxy2, "caption": "", "meta": meta,
                })
        elif is_figure_block_category(cat):
            count = len([x for x in objects if x["kind"] == "figure" and x["page_no"] == page_no]) + 1
            fig_ref = f"{safe}_page{int(page_no or 0):03d}_figure{count:03d}" if kind != "main" else f"page{int(page_no or 0):03d}_figure{count:03d}"
            img_out = fig_dir / f"{fig_ref}.png"
            xyxy2 = crop(page_img, bbox, img_out)
            if xyxy2:
                fig_id = uid(fig_id_prefix, doc_id, pdf_path.name, page_no, i)
                json_out = fig_dir / f"{fig_ref}.json"
                json_out.write_text(json.dumps(cell.get("raw"), ensure_ascii=False, indent=2), encoding="utf-8")
                objects.append({
                    "kind": "figure", "id": fig_id, "ref": fig_ref, "page_no": page_no,
                    "file_path": img_out, "json_path": json_out, "text": text,
                    "bbox": bbox, "xyxy": xyxy2, "caption": "", "meta": meta,
                })
        else:
            n_skipped += 1

    objects, captions = bind_captions(objects, captions)

    for obj in objects:
        meta = dict(obj["meta"])
        meta.update({
            "bbox": obj["bbox"],
            "xyxy": obj["xyxy"],
            "json_path": str(obj["json_path"]),
            "caption": obj["caption"],
            "text": obj["text"],
            "ingest_policy": "dotsocr_json_semantic_block_only",
        })
        if obj["kind"] == "figure":
            insert_figure(conn, obj["id"], doc_id, obj["page_no"], obj["ref"], obj["file_path"], obj["caption"], meta)
            insert_asset(conn, obj["id"], doc_id, fig_asset_type, obj["page_no"], obj["ref"], "", obj["file_path"], meta)
        elif obj["kind"] == "table":
            table_json = {
                "source": "dotsocr_json_block",
                "table_ref": obj["ref"],
                "text": obj["text"],
                "raw_block_path": str(obj["json_path"]),
                "metadata": meta,
            }
            insert_table(conn, obj["id"], doc_id, obj["page_no"], obj["ref"], obj["file_path"], table_json, meta)
            insert_asset(conn, obj["id"], doc_id, table_asset_type, obj["page_no"], "", obj["ref"], obj["file_path"], meta)

    (text_dir / "captions.json").write_text(json.dumps(captions, ensure_ascii=False, indent=2), encoding="utf-8")
    n_text_assets = write_page_text_assets(conn, doc_id, page_text_map, text_dir / "text_pages", text_asset_type, ocr_input, min_chars=80)

    return {
        "pages": len(pages),
        "text_blocks": n_text,
        "text_page_assets": n_text_assets,
        "figures": len([x for x in objects if x["kind"] == "figure"]),
        "tables": len([x for x in objects if x["kind"] == "table"]),
        "captions": len(captions),
        "skipped_blocks": n_skipped,
        "ingest_policy": "dotsocr_json_semantic_block_only",
        "document_part": kind,
    }


# -----------------------------------------------------------------------------
# Supplementary file parsing
# -----------------------------------------------------------------------------

TABLE_SUFFIXES = {".xlsx", ".xls", ".csv", ".tsv"}
TEXT_SUFFIXES = {".txt", ".md", ".rtf"}
DOCX_SUFFIXES = {".docx"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
STRUCTURE_SUFFIXES = {".sdf", ".smi", ".smiles", ".mol", ".mol2", ".inchi", ".cxsmi", ".cxsmiles"}
SEQUENCE_SUFFIXES = {".fa", ".fasta", ".faa", ".fna", ".gb", ".gbk"}
DATA_SUFFIXES = {".json", ".xml"}


def register_supplement_file(conn, doc_id: str, src: Path, out_dir: Path, asset_type: str, extra_meta: Optional[Dict[str, Any]] = None) -> Tuple[str, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    dst = out_dir / src.name
    if src.resolve() != dst.resolve():
        shutil.copy2(src, dst)
    asset_id = uid("supp_asset", doc_id, src.name, asset_type)
    meta = {"source_file": str(src), "document_part": "supplementary", **(extra_meta or {})}
    insert_asset(conn, asset_id, doc_id, asset_type, None, "", "", dst, meta)
    return asset_id, dst


def parse_excel_file(conn, doc_id: str, p: Path, out: Path) -> int:
    n = 0
    sheets = pd.read_excel(p, sheet_name=None)
    for sheet, df in sheets.items():
        safe = safe_name(str(sheet))
        csv_path = out / f"{safe_name(p.stem)}_{safe}.csv"
        df.to_csv(csv_path, index=False)
        table_id = uid("supp_table", doc_id, p.name, sheet)
        meta = {
            "source": "supplementary_excel",
            "source_file": str(p),
            "sheet_name": str(sheet),
            "saved_csv_path": str(csv_path),
            "n_rows": int(df.shape[0]),
            "n_cols": int(df.shape[1]),
            "document_part": "supplementary",
        }
        table_json = df_to_table_json(df, "supplementary_excel", meta)
        table_ref = f"{p.name}:{sheet}"
        insert_table(conn, table_id, doc_id, None, table_ref, csv_path, table_json, meta)
        insert_asset(conn, table_id, doc_id, "supplementary_table", None, "", table_ref, csv_path, meta)
        n += 1
    return n


def parse_csv_or_tsv_file(conn, doc_id: str, p: Path, out: Path) -> int:
    sep = "\t" if p.suffix.lower() == ".tsv" else None
    df = pd.read_csv(p, sep=sep)
    dst = out / p.name
    df.to_csv(dst, index=False)
    table_id = uid("supp_table", doc_id, p.name)
    meta = {
        "source": "supplementary_tsv" if p.suffix.lower() == ".tsv" else "supplementary_csv",
        "source_file": str(p),
        "saved_csv_path": str(dst),
        "n_rows": int(df.shape[0]),
        "n_cols": int(df.shape[1]),
        "document_part": "supplementary",
    }
    table_json = df_to_table_json(df, meta["source"], meta)
    insert_table(conn, table_id, doc_id, None, p.name, dst, table_json, meta)
    insert_asset(conn, table_id, doc_id, "supplementary_table", None, "", p.name, dst, meta)
    return 1


def parse_docx_file(conn, doc_id: str, p: Path, out: Path) -> int:
    _, dst = register_supplement_file(conn, doc_id, p, out, "supplementary_docx")
    doc = Document(str(p))
    paragraphs = [x.text for x in doc.paragraphs if x.text and x.text.strip()]
    text = "\n".join(paragraphs)
    block_id = uid("supp_text", doc_id, p.name)
    insert_text(conn, block_id, doc_id, None, "supplementary_docx", text, {"source_file": str(p), "saved_path": str(dst), "n_paragraphs": len(paragraphs)})

    # Also expose Word tables as raw_table, because supplementary structure/sequence tables often hide here.
    n_tables = 0
    for idx, table in enumerate(doc.tables, start=1):
        rows = []
        for r in table.rows:
            rows.append([clean(c.text) for c in r.cells])
        if not rows:
            continue
        max_cols = max(len(r) for r in rows)
        header = rows[0] if rows else []
        if len(header) != max_cols or len(set(header)) < len(header):
            header = [f"col_{i+1}" for i in range(max_cols)]
            data_rows = rows
        else:
            data_rows = rows[1:]
        normalized = []
        for r in data_rows:
            r = r + [""] * (max_cols - len(r))
            normalized.append({header[i] or f"col_{i+1}": r[i] for i in range(max_cols)})
        table_json = {
            "source": "supplementary_docx_table",
            "columns": header,
            "rows": normalized,
            "n_rows": len(normalized),
            "n_cols": max_cols,
            "metadata": {"source_file": str(p), "table_index": idx, "document_part": "supplementary"},
        }
        table_id = uid("supp_docx_table", doc_id, p.name, idx)
        table_path = out / f"{safe_name(p.stem)}_table_{idx:03d}.json"
        table_path.write_text(json.dumps(table_json, ensure_ascii=False, indent=2), encoding="utf-8")
        insert_table(conn, table_id, doc_id, None, f"{p.name}:table_{idx}", table_path, table_json, table_json["metadata"])
        insert_asset(conn, table_id, doc_id, "supplementary_table", None, "", f"{p.name}:table_{idx}", table_path, table_json["metadata"])
        n_tables += 1
    return 1 + n_tables


def parse_text_like_file(conn, doc_id: str, p: Path, out: Path, asset_type: str, section: str) -> int:
    _, dst = register_supplement_file(conn, doc_id, p, out, asset_type)
    text = read_text_safe(p, max_chars=None)
    block_id = uid("supp_text", doc_id, p.name, section)
    insert_text(conn, block_id, doc_id, None, section, text, {"source_file": str(p), "saved_path": str(dst), "asset_type": asset_type})
    return 1


def parse_supplementary_pdf(doc_id: str, pdf_path: Path, supp_work_dir: Path, conn, dots_root: str, dots_python: str, num_thread: int, dots_port: int, register_page_images: bool) -> Dict[str, Any]:
    return parse_pdf_with_dotsocr_once(
        doc_id=doc_id,
        pdf_path=pdf_path,
        work_dir=supp_work_dir,
        conn=conn,
        dots_root=dots_root,
        dots_python=dots_python,
        num_thread=num_thread,
        dots_port=dots_port,
        kind="supplementary",
        register_page_images=register_page_images,
    )


def parse_supplementary(doc_id: str, supp_dir: str, work_dir: Path, conn, dots_root: str, dots_python: str, num_thread: int, dots_port: int, parse_supp_pdf: bool = True, register_page_images: bool = True) -> Dict[str, Any]:
    stat = {
        "supplementary_items": 0,
        "supplementary_tables": 0,
        "supplementary_text_blocks": 0,
        "supplementary_pdfs_parsed": 0,
        "supplementary_assets": 0,
        "supplementary_errors": [],
    }
    if not supp_dir:
        return stat
    supp = Path(supp_dir)
    if not supp.exists():
        stat["supplementary_errors"].append({"path": str(supp), "error": "supplement_dir_not_found"})
        return stat

    out = work_dir / "supplementary"
    out.mkdir(parents=True, exist_ok=True)

    for p in sorted(supp.rglob("*")):
        if not p.is_file():
            continue
        suffix = p.suffix.lower()
        try:
            if suffix in {".xlsx", ".xls"}:
                n = parse_excel_file(conn, doc_id, p, out)
                stat["supplementary_tables"] += n
                stat["supplementary_items"] += n
            elif suffix in {".csv", ".tsv"}:
                n = parse_csv_or_tsv_file(conn, doc_id, p, out)
                stat["supplementary_tables"] += n
                stat["supplementary_items"] += n
            elif suffix in DOCX_SUFFIXES:
                n = parse_docx_file(conn, doc_id, p, out)
                stat["supplementary_items"] += n
                stat["supplementary_text_blocks"] += 1
            elif suffix in TEXT_SUFFIXES:
                parse_text_like_file(conn, doc_id, p, out, "supplementary_text_file", f"supplementary_{suffix[1:]}")
                stat["supplementary_text_blocks"] += 1
                stat["supplementary_items"] += 1
            elif suffix in STRUCTURE_SUFFIXES:
                parse_text_like_file(conn, doc_id, p, out, "supplementary_structure_file", f"supplementary_structure_{suffix[1:].replace('.', '')}")
                stat["supplementary_assets"] += 1
                stat["supplementary_items"] += 1
            elif suffix in SEQUENCE_SUFFIXES:
                parse_text_like_file(conn, doc_id, p, out, "supplementary_sequence_file", f"supplementary_sequence_{suffix[1:].replace('.', '')}")
                stat["supplementary_assets"] += 1
                stat["supplementary_items"] += 1
            elif suffix in DATA_SUFFIXES:
                register_supplement_file(conn, doc_id, p, out, "supplementary_data_file")
                stat["supplementary_assets"] += 1
                stat["supplementary_items"] += 1
            elif suffix == ".pdf":
                if parse_supp_pdf and dots_root and dots_python:
                    pdf_stat = parse_supplementary_pdf(doc_id, p, out, conn, dots_root, dots_python, num_thread, dots_port, register_page_images)
                    stat["supplementary_pdfs_parsed"] += 1
                    stat["supplementary_items"] += 1 + int(pdf_stat.get("figures", 0)) + int(pdf_stat.get("tables", 0))
                else:
                    register_supplement_file(conn, doc_id, p, out, "supplementary_pdf", {"parsed": False})
                    stat["supplementary_items"] += 1
            elif suffix in IMAGE_SUFFIXES:
                register_supplement_file(conn, doc_id, p, out, "supplementary_image")
                stat["supplementary_items"] += 1
                stat["supplementary_assets"] += 1
            else:
                register_supplement_file(conn, doc_id, p, out, "supplementary_file")
                stat["supplementary_items"] += 1
                stat["supplementary_assets"] += 1
        except Exception as e:
            stat["supplementary_errors"].append({"path": str(p), "error": str(e)})
    return stat


# -----------------------------------------------------------------------------
# Validation and main
# -----------------------------------------------------------------------------

def validate_parse_outputs(conn, doc_id: str) -> Dict[str, Any]:
    counts = {}
    for table in ["raw_asset", "raw_text_block", "raw_figure", "raw_table"]:
        counts[table] = int(conn.execute(f"SELECT COUNT(*) AS n FROM {table} WHERE doc_id=?", (doc_id,)).fetchone()["n"])

    missing_asset_files = []
    for row in conn.execute("SELECT asset_id, file_path FROM raw_asset WHERE doc_id=?", (doc_id,)).fetchall():
        fp = row["file_path"] or ""
        if fp and not Path(fp).exists():
            missing_asset_files.append({"asset_id": row["asset_id"], "file_path": fp})

    page_assets = int(conn.execute(
        "SELECT COUNT(*) AS n FROM raw_asset WHERE doc_id=? AND asset_type IN ('page_image','supplementary_page_image')",
        (doc_id,),
    ).fetchone()["n"])

    warnings = []
    if counts["raw_text_block"] == 0:
        warnings.append("no_raw_text_block")
    if page_assets == 0:
        warnings.append("no_page_image_asset")
    if missing_asset_files:
        warnings.append("missing_raw_asset_files")

    return {
        "counts": counts,
        "page_image_assets": page_assets,
        "missing_asset_files": missing_asset_files[:50],
        "num_missing_asset_files": len(missing_asset_files),
        "warnings": warnings,
        "ok": not missing_asset_files and counts["raw_text_block"] > 0 and page_assets > 0,
    }


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--doc-id", required=True)
    ap.add_argument("--dots-root", default="/root/autodl-tmp/dots.ocr")
    ap.add_argument("--dots-python", default="/root/autodl-tmp/miniconda3/envs/dots_mocr/bin/python", help="dots_mocr的环境python")
    ap.add_argument("--num-thread", type=int, default=16)
    ap.add_argument("--dots-port", type=int, default=8001)
    ap.add_argument("--skip-supp-pdf-ocr", action="store_true", help="Register supplementary PDFs without OCR parsing.")
    ap.add_argument("--no-page-assets", action="store_true", help="Render page images but do not register them as raw_asset.")
    ap.add_argument("--overwrite", action="store_true", help="Delete existing raw parser outputs for this doc_id before parsing.")
    args = ap.parse_args()

    conn = get_conn()
    row = conn.execute(
        "SELECT doc_id, source_pdf_path, supplement_dir FROM raw_document WHERE doc_id=?",
        (args.doc_id,),
    ).fetchone()
    if not row:
        raise SystemExit(f"doc_id not found: {args.doc_id}")

    pdf_path = Path(row["source_pdf_path"])
    if not pdf_path.exists():
        update_document_status(conn, args.doc_id, "parse_failed", {"parse_error": f"PDF not found: {pdf_path}"})
        conn.commit()
        raise SystemExit(f"PDF not found: {pdf_path}")

    if args.overwrite:
        clear_existing_raw_outputs(conn, args.doc_id)
        conn.commit()

    work_dir = Path("data/work") / args.doc_id
    work_dir.mkdir(parents=True, exist_ok=True)

    status = "parsed_dotsocr"
    parse_error = ""
    stat: Dict[str, Any] = {}
    supp_stat: Dict[str, Any] = {}
    validation: Dict[str, Any] = {}

    try:
        stat = parse_pdf_with_dotsocr_once(
            doc_id=args.doc_id,
            pdf_path=pdf_path,
            work_dir=work_dir,
            conn=conn,
            dots_root=args.dots_root,
            dots_python=args.dots_python,
            num_thread=args.num_thread,
            dots_port=args.dots_port,
            kind="main",
            register_page_images=not args.no_page_assets,
        )
        supp_stat = parse_supplementary(
            doc_id=args.doc_id,
            supp_dir=row["supplement_dir"],
            work_dir=work_dir,
            conn=conn,
            dots_root=args.dots_root,
            dots_python=args.dots_python,
            num_thread=args.num_thread,
            dots_port=args.dots_port,
            parse_supp_pdf=not args.skip_supp_pdf_ocr,
            register_page_images=not args.no_page_assets,
        )
        validation = validate_parse_outputs(conn, args.doc_id)
        if validation.get("warnings"):
            status = "parsed_with_warnings"
        update_document_status(
            conn,
            args.doc_id,
            status,
            {
                "parse_status": status,
                "parse_stat": stat,
                "supplementary_parse_stat": supp_stat,
                "parse_validation": validation,
            },
        )
        conn.commit()
    except Exception as e:
        parse_error = str(e)
        status = "parse_failed"
        update_document_status(conn, args.doc_id, status, {"parse_error": parse_error})
        conn.commit()
        conn.close()
        raise

    conn.close()
    print(json.dumps({
        "doc_id": args.doc_id,
        "status": status,
        "work_dir": str(work_dir),
        "main_pdf": stat,
        "supplementary": supp_stat,
        "validation": validation,
        "error": parse_error,
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
